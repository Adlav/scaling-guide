from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook


address = ""
trust =""
date =""
dear =""
intro =""
appt =""
outro =""
signature =""


book =  load_workbook('Sample.xlsx')
sheet1 = book.active
doc = Document()
data = []
for i in sheet1:
    for j in i:
        data.append(j.value)

print data

header = doc.sections[0]
header.add_paragraph('Hello')
doc.add_picture('Repselfie-Logo-B.png', width=None, height=None)
doc.add_paragraph('Dear ' + str(data[2]))
doc.add_paragraph('')
doc.save('sample.docx')
